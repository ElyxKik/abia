
import sys
import os
import json
import docx

def extract_text_from_word(docx_path):
    """Extract text and metadata from a Word document."""
    if not os.path.exists(docx_path):
        print(f"Error: File {docx_path} does not exist.", file=sys.stderr)
        sys.exit(1)
    
    try:
        doc = docx.Document(docx_path)
        
        # Extract text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract metadata
        metadata = {
            "fileName": os.path.basename(docx_path),
            "paragraphCount": len(doc.paragraphs),
            "sectionCount": len(doc.sections),
            "tableCount": len(doc.tables)
        }
        
        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)
        
        result = {
            "metadata": metadata,
            "text": "\n".join(full_text),
            "tables": tables
        }
        
        return result
    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python word_text_extractor.py <docx_file_path>", file=sys.stderr)
        sys.exit(1)
    
    docx_path = sys.argv[1]
    result = extract_text_from_word(docx_path)
    print(json.dumps(result))
