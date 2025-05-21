#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Document Extractor
-----------------
Script Python pour l'extraction de texte à partir de documents PDF et Word.
Utilisé par l'agent Document de l'application ABIA.
"""

import sys
import json
import os
from pathlib import Path

def extract_text_from_pdf(file_path):
    """
    Extrait le texte d'un fichier PDF.
    
    Args:
        file_path (str): Chemin vers le fichier PDF
    
    Returns:
        str: Texte extrait du PDF
    """
    try:
        import PyPDF2
        
        # Vérifier si le fichier existe
        if not Path(file_path).exists():
            return json.dumps({"error": f"Le fichier {file_path} n'existe pas."})
        
        # Extraire le texte du PDF
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            metadata = {
                "page_count": num_pages,
                "title": reader.metadata.get('/Title', ''),
                "author": reader.metadata.get('/Author', ''),
                "subject": reader.metadata.get('/Subject', ''),
                "creator": reader.metadata.get('/Creator', ''),
                "producer": reader.metadata.get('/Producer', '')
            }
            
            # Extraire le texte de chaque page
            pages = []
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                pages.append(page_text)
                text += page_text + "\n\n"
            
            result = {
                "metadata": metadata,
                "text": text,
                "pages": pages
            }
            
            return json.dumps(result)
    
    except Exception as e:
        return json.dumps({"error": str(e)})

def extract_text_from_docx(file_path):
    """
    Extrait le texte d'un fichier Word (DOCX).
    
    Args:
        file_path (str): Chemin vers le fichier DOCX
    
    Returns:
        str: Texte extrait du document Word
    """
    try:
        import docx
        
        # Vérifier si le fichier existe
        if not Path(file_path).exists():
            return json.dumps({"error": f"Le fichier {file_path} n'existe pas."})
        
        # Extraire le texte du document Word
        doc = docx.Document(file_path)
        
        # Extraire les métadonnées
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "section_count": len(doc.sections),
            "table_count": len(doc.tables)
        }
        
        # Extraire le texte de chaque paragraphe
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        
        # Extraire le texte des tableaux
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)
        
        result = {
            "metadata": metadata,
            "text": "\n".join(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables
        }
        
        return json.dumps(result)
    
    except Exception as e:
        return json.dumps({"error": str(e)})

if __name__ == "__main__":
    # Vérifier si un chemin de fichier a été fourni
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Aucun chemin de fichier fourni."}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        print(extract_text_from_pdf(file_path))
    elif file_ext in ['.docx', '.doc']:
        print(extract_text_from_docx(file_path))
    else:
        print(json.dumps({"error": f"Format de fichier non pris en charge: {file_ext}"}))
